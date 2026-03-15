import os
import sys
import subprocess
import re

def install_and_import(package):
    try:
        import docx
    except ImportError:
        print(f"Installing {package}...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", package])
        import docx

install_and_import('python-docx')
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def __add_run(p, text, is_bold=False, is_italic=False, size=10, is_small_caps=False):
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = is_bold
    run.font.italic = is_italic
    run.font.small_caps = is_small_caps
    return run

def parse_markdown_to_docx(md_path, docx_path):
    doc = Document()
    
    # Page Layout
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(1.69)
    section.left_margin = Inches(0.56)
    section.right_margin = Inches(0.56)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Split into frontmatter before --- and content after ---
    frontmatter = []
    content = []
    is_frontmatter = True
    for line in lines:
        stripped = line.strip()
        if stripped == '---':
            if is_frontmatter:
                is_frontmatter = False
                continue
        if is_frontmatter:
            frontmatter.append(stripped)
        else:
            content.append(stripped)
            
    # Parse frontmatter
    title_text = ""
    authors = []
    abstract = ""
    keywords = ""
    
    for line in frontmatter:
        if line.startswith('# '):
             title_text = line[2:].strip()
        elif line.startswith('**Abstract**—'):
             abstract = line[13:].strip()
        elif line.startswith('**Keywords**—'):
             keywords = line[13:].strip()
        elif line: # non-empty line
             authors.append(line)

    # Add Title
    title_p = doc.add_paragraph(title_text)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(12)
    for run in title_p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(24)

    # Add Authors
    # (combine multiple lines into a single centered block)
    authors_text = '\n'.join([a.replace('**', '').replace('*', '') for a in authors])  # simple strip
    authors_p = doc.add_paragraph(authors_text)
    authors_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors_p.paragraph_format.space_after = Pt(24)
    for run in authors_p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

    # Abstract
    if abstract:
        abs_p = doc.add_paragraph()
        abs_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        abs_p.paragraph_format.space_after = Pt(6)
        __add_run(abs_p, 'Abstract— ', is_bold=True, size=9)
        __add_run(abs_p, abstract, size=9)
    
    # Keywords
    if keywords:
        kw_p = doc.add_paragraph()
        kw_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        kw_p.paragraph_format.space_after = Pt(18)
        __add_run(kw_p, 'Keywords— ', is_bold=True, is_italic=True, size=9)
        __add_run(kw_p, keywords, size=9)

    # Two column section
    new_section = doc.add_section(0) # Continuous section break
    sectPr = new_section._sectPr
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '708') # 0.5 inches spacing between columns
    sectPr.append(cols)

    # Content parser
    in_table = False
    table_data = []

    def process_body_text(p, text, size=10):
        # Very simple bold/italic parser: **bold**, *italic*
        # We'll split on ** and * using regex
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                __add_run(p, part[2:-2], is_bold=True, size=size)
            elif part.startswith('*') and part.endswith('*'):
                __add_run(p, part[1:-1], is_italic=True, size=size)
            elif part:
                __add_run(p, part, size=size)

    def write_table():
        if not table_data: return
        # Create table
        rows = len(table_data)
        cols = max(len(row) for row in table_data) if table_data else 0
        if cols == 0: return
        
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        
        # Center the table slightly if needed (Word handles this within column automatically)
        for r_idx, row in enumerate(table_data):
            for c_idx, cell_text in enumerate(row):
                if c_idx < len(table.rows[r_idx].cells):
                    cell = table.rows[r_idx].cells[c_idx]
                    cell_p = cell.paragraphs[0]
                    process_body_text(cell_p, cell_text.strip(), size=9)
        # add space after table
        table_space_p = doc.add_paragraph()
        table_space_p.paragraph_format.space_after = Pt(6)
        table_data.clear()

    for line in content:
        if line.startswith('|'):
            in_table = True
            # Parse row
            if '---' in line and set(line.strip('|').replace('-', '').replace(':', '').replace(' ', '')) == set():
                # skip separator row
                continue
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            table_data.append(cells)
            continue
        elif in_table:
            in_table = False
            write_table()
            
        if not line:
            # We don't add empty paragraphs, Word handles inter-paragraph spacing now!
            continue

        if line.startswith('### '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(12) # Huge gap before heading
            p.paragraph_format.space_after = Pt(6) # Medium gap after heading
            __add_run(p, line[4:].upper(), is_bold=True, size=10, is_small_caps=True)
            continue
            
        if line.startswith('![') and '](' in line:
            # Image
            # ![Architecture...](/path/to/img.png)
            m = re.match(r'!\[(.*?)\]\((.*?)\)', line)
            if m:
                caption = m.group(1)
                img_path = m.group(2)
                try:
                    if os.path.exists(img_path):
                        img_para = doc.add_paragraph()
                        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        img_para.paragraph_format.space_before = Pt(12)
                        img_para.paragraph_format.space_after = Pt(4) # less space since caption is below
                        run = img_para.add_run()
                        run.add_picture(img_path, width=Inches(3.2)) # Fits into column width easily
                except Exception as e:
                    print(f"Could not add image {img_path}: {e}")
            continue

        # Normal paragraph (check if it's a table caption e.g. *Table 1...* or fig caption *Fig. 1...*)
        if line.startswith('*Fig.') or line.startswith('*Table'):
             p = doc.add_paragraph()
             p.alignment = WD_ALIGN_PARAGRAPH.CENTER
             p.paragraph_format.space_after = Pt(12) # Padding after figures
             process_body_text(p, line, size=8)
             # Apply small caps to runs if it's Table caption
             if 'Table' in line:
                for run in p.runs:
                     run.font.small_caps = True
             continue

        # References
        if line.startswith('['):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_after = Pt(4)
            process_body_text(p, line, size=8)
            continue

        # Normal body paragraph
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6) # standard gap
        process_body_text(p, line, size=10)
        p.paragraph_format.first_line_indent = Inches(0.15)

    if table_data:
        write_table()

    doc.save(docx_path)
    print(f"Document updated: {docx_path}")

if __name__ == "__main__":
    md_file = "/Users/apple/Desktop/compiler design /NSIR_Research_Paper.md"
    docx_file = "/Users/apple/Desktop/NS-IR_Compiler_Research_Paper.docx"
    print(f"Parsing {md_file} to {docx_file}")
    if os.path.exists(md_file):
        parse_markdown_to_docx(md_file, docx_file)
    else:
        print(f"Could not find markdown file at: {md_file}")
